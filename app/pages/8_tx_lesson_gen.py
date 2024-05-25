from typing import List, Optional, Dict
import time
import streamlit as st
from openai import OpenAI
import os
import io
import sys
import re
import json
from utils_8 import assistant_ids, examples
from langchain_openai import ChatOpenAI
from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import JsonOutputParser
from langchain_core.pydantic_v1 import BaseModel, Field
from docx import Document

client = OpenAI(api_key=os.getenv('TX_OPENAI_API_KEY'))

class LearningTargets(BaseModel):
    lt1: str = Field(..., 
                     description="First learning target for a lesson to align with the standards, starts with 'I can' and is a building block for the next learning target.", 
                     example="I can partition and label equal-sized parts with unit fractions.")
    lt2: str = Field(..., 
                     description="Second learning target for a lesson to align with the standards, starts with 'I can' and builds on the first learning target.", 
                     example="I can partition rectangles by drawing and continue to practice naming the parts with the unit fractions.")

def lts_prompt(parser):
    prompt = PromptTemplate(
        template = """You are a helpful assistant to write learning targets for a lesson plan based on standards for Texas math curriculum.  
        You are provided with the standard <standard>{standard}</standard> and its description <desc>{standard_description}</desc> and a few examples, author two learning targets for the lesson.
        The language should be clear and concise, age appropriate for {grade}.
          <examples>
          {examples}
          </examples>
        and output them in the following format {format_instructions}
        """,
        input_variables = ['grade', "standard", "standard_description", "examples"],
        partial_variables={"format_instructions": parser.get_format_instructions()},
    )
    return prompt
@st.cache_data()
def run_learning_targets(standard, standard_description, grade, asst_lg_id):
    parser = JsonOutputParser(pydantic_object=LearningTargets)
    prompt = lts_prompt(parser)
    content = prompt.invoke({
        "grade": grade,
        "standard": standard, 
        "standard_description": standard_description,
        "examples": examples[grade],
    })
    print(content)
    try:
        thread = client.beta.threads.create(
                messages = [{'role': 'user','content': content.text}],
        )
        run = client.beta.threads.runs.create(
                thread_id = thread.id,
                assistant_id = asst_lg_id
            )
        check_run_status(run, thread.id, 'learning_targets')
        return thread, run
    except Exception as e:
        st.write(f'Error: {e}')
        sys.exit(1)
        
def check_run_status(run, thread_id, fu_key, sleep=5):
    while run.status != 'completed':
        run = client.beta.threads.runs.retrieve(
                thread_id=thread_id,
                run_id=run.id
            )
        time.sleep(sleep)
    return run


class ActivityQuestion(BaseModel):
    question: str = Field(..., title="Question", description='Question text')
    question_type: str = Field(..., title="Question Type", description='Type of question, e.g. multiple choice, multiple select, written response')
    student_response: str = Field(..., title="Student Response", description='Detailed step-by-step response to the task and avoid generic answers; if the activity is open-ended, provide sample response')

class Activity(BaseModel):
    title: str = Field(..., title="Activity Title", description='Title of the activity')
    addressing: str = Field(..., title="Addressing", description='Standard name that is being addressed for the lesson')
    purpose: str = Field(..., title="Purpose", description='Purpose of the activity, describe the purpose of the activity in up to two sentences. Start with "The purpose of this activity is to..." and align with the provided learning target.')
    narrative: str = Field(..., title="Activity Narrative", description='Narrative of the activity, start with "In this activity," state how students are interacting with the activity.')
    required_materials: Dict[str, List[str]] = Field(default=None, title="Required Materials", description='Optional: Materials needed for the activity. Provide a dictionary with keys "Materials to Gather" and "Materials to Copy", each with a list of items.')
    launch_instructions: List[str] = Field(..., title="Launch", description='A list.', example=['Groups of 2.'])
    student_facing_task_statement: str = Field(..., title="Student-facing Task Statement", description='Task statement for the activity. Provide the specific numbers of the task that students will work on either individually or in a small group. The statement shall be specific while the language needs to be simple and clear.')
    questions: List[ActivityQuestion] = Field(default=None, title="Questions", description='List of return a list of 3 to 8 practice questions, multiple choice or multiple select or written response, that students will work on to achieve the learning target.')
    learning_target: str = Field(..., title="Learning Target", description='Learning target that the activity is aligned with. This is given.')
    activity_synthesis: List[str] = Field(..., title="Activity Synthesis", description='List of actions or questions that students will take to synthesize their learning.')
    advancing_student_thinking: str = Field(..., title="Advancing Student Thinking", description='Anticipate possible misconceptions and how to address them in advancing questions.')
    duration: int = Field(..., title="Duration", description='Duration of the activity in minutes, between 15 and 20 minutes.')

class Warmup(BaseModel):
    title: str = Field(..., title="Warm-up Title", description='Title of the warm-up activity')
    purpose: str = Field(..., title="Purpose",
                        description='State the purpose of the warmup activity, specifically how it prepares students for the day\'s lesson or strengthens their number sense or procedural fluency.')
    standards: str = Field(..., title="Standards", description="Standard code that the warm-up activity aligns with.")
    instructional_routines: str = Field(..., title="Instructional Routines", 
                                        description="Instructional routine to be applied in the warm-up activity. Shall be one of the provided instructional routines.")
    activity_narrative: str = Field(..., title="Activity Narrative", 
                                    description="Narrative of how students interact with the activity, start with 'In this activity,'. The activity shall take a problem-based learning approach, and the purpose is to prepare students to start the first activity.")
    launch: str = Field(..., title="Launch", description="Instructions for launching the activity.")
    student_facing_task_statement: str = Field(..., title="Student-facing Task Statement", 
                                               description="Task statement for students,apply the instructional routine selected and write a math problem or problems. make sure it is age-appropriate for the grade and engaging for the students")
    student_response: str = Field(..., title="Student Response", 
                                  description="Detailed step-by-step response to the task, avoiding generic answers; provide sample response if open-ended.")
    activity_synthesis: str = Field(..., title="Activity Synthesis", 
                                    description="Questions or actions to synthesize the learning, preparing students for Activity 1 & 2 in the lesson.")
    duration: int = Field(..., title="Duration", description="Duration of the warm-up activity in minutes, between 5 and 10 minutes.")

def prompt_activity_1(pydantic_object = Activity):
    parser = JsonOutputParser(pydantic_object = pydantic_object)
    prompt = PromptTemplate(
        template = """Write a math lesson that aligns with 
                    Texas standards <standards>{standards}</standards> for {grade}.
                    The lesson shall address the learning targets <learning_targets>{learning_targets}</learning_targets>.
                        
                    Let's start with two activities that are the heart of the mathematical experience and make up the majority of the time spent in class.
                    Each activity can serve one or more of many purposes: 
                    - Provide experience with a new context.
                    - Introduce a new concept and associated language.
                    - Introduce a new representation.
                    - Formalize a definition of a term for an idea previously encountered informally.
                    - Identify and resolve common mistakes and misconceptions that people make.
                    - Practice using mathematical language.
                    - Work toward mastery of a concept or procedure.
                    - Provide an opportunity to apply mathematics to a modeling or other application problem.
                    
                    A typical activity has four phases:
                    1. A launch (Engagement Phase)
                        - Goal: Introduce the context and problem of the lesson.
                        - Focus: Ensure students understand the problem, not necessarily how to solve it.
                        - Techniques: Use grouping suggestions, context setup, and problem clarification.
                    2. Student work time (Exploration Phase)
                        - Goal: Allow students to engage with and explore the mathematical concepts.
                        - Focus: Facilitate individual, partnered, or small group work.
                        - Techniques: Encourage problem-solving, critical thinking, and peer collaboration.
                    3. Activity synthesis (Reflection Phase)
                        - Goal: Help students synthesize and internalize their learning.
                        - Focus: Connect new learning to prior knowledge and the unit's big picture.
                        - Techniques: Use questioning, journaling, graphic organizers, or concept maps.
                    4. Cool-Down (Assessment Phase)
                        - Goal: Conduct a brief formative assessment of students’ understanding.
                        - Focus: Evaluate if students grasped the lesson’s key concepts.
                        - Techniques: Assign a short, independent task to assess learning.  

                    And each activity routines embed structures within the tasks of the lessons that allow students to engage in the content, 
                    and collaborate in ways that support the development of student thinking and precision with language. 
                    MLRs are written into each lesson, either as an embedded structure of a lesson activity in which all students engage, 
                    or as a suggested optional support specifically for English learners.
                    Each activity includes a synthesis that provides an opportunity for students to discuss key mathematical ideas of the activity/lesson 
                    and incorporate their new insights into their big-picture understanding.

                    # Additional Notes for Lesson Authoring
                    - Standards Alignment: Ensure activities align with grade-appropriate standards.
                    - Student Engagement: Craft activities that are challenging, relevant, and engaging, catering to the developmental stage of high school students.
                    - Feedback and Adaptation: Provide opportunities for teachers to use student responses and feedback to refine and adapt lessons.
                    - Diverse Learning Styles: Include a range of instructional strategies to address different learning preferences and abilities.
                    - Language Considerations
                        Simplicity and Clarity: Use simple, clear language. Avoid complex terms that might confuse them.
                        Language needs to be age-appropriate for {grade}, removing puffy language.
                    - Apply Problem-based approach. the definition of Problem-based learning is to present students with scenarios or problems that require students to either
                    a) apply the skills needed to meet the activity's "i can" statement, or 
                    b) explore and discover the concepts stated in the activity's "i can" statement. 
                    The important part is that we are not explicitly explaining ideas or teaching concepts,
                    but instead provide a space for students to take ownership of the learning, 
                    relying on their own knowledge and discussions with their peers to meet the lesson's goals.
                    - Activities are designed to be engaging, apply a problem-based learning approach that encourages collaboration among students and individual critical thinking.
                    
                Output the first Activity that address the first learning target ```{first_learning_target}```
                in the following format {format_instructions}""",
        input_variables = ["standards", "learning_targets", "grade", "first_learning_target", "alignment_note"],
        partial_variables={"format_instructions": parser.get_format_instructions()},
    )   
    return prompt
def prompt_activity_2(pydantic_object = Activity):
    parser = JsonOutputParser(pydantic_object = pydantic_object)
    prompt = PromptTemplate(
        template = """Following the guidelines, generate the second activity for the same lesson.    
                Recall, the lesson needs to align with
                standards <standards>{standards}</standards> for {grade}.
                This activity shall build upon the first activity and specifically address the second learning target ```{second_learning_target}```
                in the following format {format_instructions}""",
        input_variables = ["second_learning_target", "standards", "grade"],
        partial_variables={"format_instructions": parser.get_format_instructions()},
    )   
    return prompt

def prompt_warmup(pydantic_object = Warmup):
    parser = JsonOutputParser(pydantic_object = pydantic_object)
    prompt = PromptTemplate(
        template = """Write a warm-up activity for the math lesson that aligns with
                Texas standards <standards>{standards}</standards> for {grade}; the activity shall take 5-10 minutes in class. 

                # Warm-up Activity Design Principles:
                - Purpose: Engage students in the lesson's mathematics, incorporating their experiences and knowledge. Prepare students for the day's lesson while enhancing number sense and procedural fluency. They place value on students' voices as they communicate their developing ideas, ask questions, justify their responses, and critique the reasoning of others.
                - Types:
                    - Preparation for the day's lesson, strengthening number sense/procedural fluency.
                    -  **Warm-up routines**.
                        - Number Talk: The sequence of problems in a Number Talk encourages students to look for structure and use repeated reasoning to evaluate expressions and develop computational fluency (MP7 and MP8). As students share their strategies, they make connections and build on one another’s ideas, developing conceptual understanding.
                        - Notice and Wonder: Notice and Wonder invites all students into a mathematical task with two low-stakes prompts: “What do you notice? What do you wonder?” By thinking about things they notice and wonder, students gain entry into the context and might have their curiosity piqued. Students learn to make sense of problems (MP1) by taking steps to become familiar with a context and the mathematics that might be involved. Note: Notice and Wonder and I Notice/I Wonder are trademarks of NCTM and the Math Forum and are used in these materials with permission.
                        - True or False?: The True or False routine structure encourages students to reason about numerical expressions and equations using base-ten structure, meaning and properties of operations, and the meaning of the equal sign. Often, students can determine whether an equation or inequality is true or false without doing any direct computation (MP7).
                        - Which One Doesn’t Belong?: Which One Doesn’t Belong fosters a need for students to identify defining attributes and use language precisely in order to compare and contrast a carefully chosen group of geometric figures, images, or other mathematical representations (MP3 and MP6).
                - Instructional Routines: Use routines to strengthen listening and speaking skills in mathematics.
                - Implementation Tips: Streamline the process for efficiency; consider hand signals for student responses to maintain focus. Manage time and student responses effectively.

                 - Language Considerations
                    Simplicity and Clarity: Use simple, clear language. Avoid complex terms that might confuse them.
                    Interactive Language: Encourage interaction through language. Ask questions, prompt discussions, and encourage children to express their thoughts and understanding of mathematical concepts.
                - Math Activities Considerations
                    Apply Problem-based approach. the definition of Problem-based learning is to present students with scenarios or problems that require students to either
                    a) apply the skills needed to meet the activity's "i can" statement, or 
                    b) explore and discover the concepts stated in the activity's "i can" statement. 
                    The important part is that we are not explicitly explaining ideas or teaching concepts,
                    but instead provide a space for students to take ownership of the learning, 
                    relying on their own knowledge and discussions with their peers to meet the lesson's goals.

                Output the warm-up activity in the following format {format_instructions}""",
        input_variables = ["standards", "grade"],
        partial_variables={"format_instructions": parser.get_format_instructions()},
    )
    return prompt

class LessonSynthesis(BaseModel):
    lesson_synthesis: str = Field(..., title="Lesson Synthesis", 
                                  description="return a paragraph that summarizes the key takeaways from the lesson, and how it connects to the learning targets")
    duration: int = Field(..., title="Duration", description="Duration of the lesson synthesis in minutes, between 5 and 10 minutes.")

def prompt_lesson_synthesis(pydantic_object = LessonSynthesis):
    parser = JsonOutputParser(pydantic_object = pydantic_object)
    prompt = PromptTemplate(
        template = """Building on Warmup, Activity 1, and Activity 2, write a lesson synthesis that shall take 5-10 mins during class. 
                After the activities for the day are done, students should take time to synthesize what they have learned. 
                The lesson synthesis assists the teacher with ways to help students incorporate new insights gained during the activities into their big-picture understanding. 
                Teachers can use this time in any number of ways, including posing questions verbally and calling on volunteers to respond, 
                asking students to respond to prompts in a written journal, 
                asking students to add on to a graphic organizer or concept map, 
                or adding a new component to a persistent display like a word wall.

                The lesson synthesis shall be aligned with learning targets <learning_targets>{learning_targets}</learning_targets>;
                especially with the second learning target ```{second_learning_target}```. 

                Make sure to use engaging language but concise sentences to get the point across.

                Output the lesson synthesis in the following format {format_instructions}""",
        input_variables = ['lesson_title', "unit_title", "standards", "grade"],
        partial_variables={"format_instructions": parser.get_format_instructions()},
    )
    return prompt

class LessonPlan(BaseModel):
    lesson_title: str = Field(..., title="Lesson Title")
    standards: str = Field(..., title="Addressing", description='Standards that the lesson aligns with')
    learning_goals_teacher: str = Field(..., title="Learning Goals - Teacher", 
                                        description='''Describe, for teachers, the mathematical and pedagogical goals of the lesson in imperative language.
                                        reframe the i can statements in learning targets into two purpose statements, and how those takeaways from each activity will help students meet the learning targets.
                                        Return in a list format, each item is a sentence.''')
    learning_goals_student: str = Field(..., title="Learning Goals - Student", 
                                        description='One short sentence. Start with the word "Let\'s."', 
                                        example=['Let’s use ratios to work with how fast things move.', 
                                                 'Let’s see how exponents show repeated multiplication.',
                                                 'Let’s explore cylinders and their volumes.',])
    lesson_purpose: str = Field(..., title="Lesson Purpose", 
                                        description='''state the goal of this lesson in one sentence, start with 'The purpose of this lesson is for students to...''')
    narrative: str = Field(..., title="Narrative", description='''Narrative for the lesson, start with 'In this lesson,' stating how the lesson is presented to the students, and the learning goal of the lesson.''')
    learning_targets: str = Field(..., title="Learning Targets", description='Start with "Students will be able to..." that paraphrase the provided learning targets of the lesson.')
    required_materials: Optional[Dict[str, List[str]]] = Field(default=None, title="Required Materials", 
                                                     description='Any Materials needed for the lesson: activities, warmup, etc. Provide a dictionary with keys "Materials to Gather" and "Materials to Copy", each with a list of items.')
    lesson_timeline: Dict[str, int] = Field(..., title="Lesson Timeline", description='Timeline for the lesson, provide a dictionary with keys "Warm-up" "Activity 1" "Activity 2" "Lesson Synthesis", and "Cool-down", each with a duration in minutes; and the sum of all durations should be 60 mins.')
    teacher_reflection_questions: str = Field(..., title="Teacher Reflection Questions", 
                                              description='RA few questions aimed for teachers to reflect on if the students met the learning targets.')
    student_reflection_questions: str = Field(..., title="Student Reflection Questions", 
                                            description='A few questions aimed for students to reflect on if they met the learning targets.')

def prompt_lesson_plan(pydantic_object = LessonPlan):
    parser = JsonOutputParser(pydantic_object = pydantic_object)
    prompt = PromptTemplate(
        template = """Finally, let's complete the other components of the lesson, 
                using the content generated in the previous steps.
                Recall the standards are <standards>{standards}</standards> for {grade}.
                The lesson is designed to help students meet the learning targets <learning_targets>{learning_targets}</learning_targets>,
                especially the second learning target ```{second_learning_target}```.

                A typical lesson has four parts, and total duration should be 60 minutes in total:
                a warm-up, takes 5-10 minutes
                activity 1, takes 15-20 minutes
                activity 2, takes 15-20 minutes
                lesson synthesis, takes 5-10 minutes
                cool-down, takes 5 minutes
                Use the duration estimate for Lesson Timeline.
                
                Keep the verbosity level low, concise and to the point.
                Output the lesson plan in the following format {format_instructions}""",
        input_variables = ['standards', 'grade', 'learning_targets', 'second_learning_target'],
        partial_variables={"format_instructions": parser.get_format_instructions()},
    )
    return prompt

class CoolDown(BaseModel):
    title: str = Field(..., title="Cool-down Title", description='Title of the cool-down activity')
    student_facing_task_statement: str = Field(..., title="Student-facing Task Statement", 
                                               description="Task statement for students, make sure it is age-appropriate for the grade and engaging for the students")
    questions: List[ActivityQuestion] = Field(default=None, title="Questions", description='List of return a list of 3 to 5 practice questions, multiple choice or multiple select or written response, that students will work on to achieve the learning target.')
    responding_to_student_thinking: str = Field(..., title="Responding to Student Thinking", 
                                    description="Anticipate possible misconceptions and how to address them in advancing questions.")
    duration: int = 5
def prompt_cool_down(pydantic_object = CoolDown):
    parser = JsonOutputParser(pydantic_object = pydantic_object)
    prompt = PromptTemplate(
        template = """Write a 5-min cool-down activity lesson to assess if the students understood the main concepts of the lesson;
                especially if they met th second learning target ```{second_learning_target}```.

                The cool-down (also known as an exit slip or exit ticket) is to be given to students at the end of the lesson. 
                This activity serves as a brief check-in to determine whether students understood the main concepts of this lesson. 
                Teachers can use this as a formative assessment to plan further instruction.

                Guidance for unfinished learning, evidenced by the cool-down, is provided in two categories: next-day support and prior-unit support. 
                This guidance is meant to provide teachers ways in which to continue grade-level content while also giving students the additional support they may need.

                Make sure the language is age-appropriate for {grade}, activities are designed to be engaging and encourage critical thinking.

                Output the cool-down activity in the following format {format_instructions}""",
        input_variables = ['standards', 'grade', 'second_learning_target'],
        partial_variables={"format_instructions": parser.get_format_instructions()},
    )
    return prompt

def len_message(thread_id):
    messages = client.beta.threads.messages.list(
                    thread_id=thread_id
                ) 
    return len(messages.data)




def create_all_prompts(row):
    p1 = prompt_activity_1()
    p2 = prompt_activity_2()
    wu = prompt_warmup()
    ls = prompt_lesson_synthesis()
    cd = prompt_cool_down()
    lp = prompt_lesson_plan()

    p1_out = p1.invoke({
        'grade': row['grade'],
        'first_learning_target': row['learning_targets'][0],
        'standards': row['standards'], 
        'learning_targets': row['learning_targets_str'], 
    })
    p2_out = p2.invoke({
        'second_learning_target': row['learning_targets'][1],
        'standards': row['standards'], 
        'grade': row['grade']
    })
    wu_out = wu.invoke({
        'standards': row['standards'], 
        'grade': row['grade']
    })
    ls_out = ls.invoke({
        'standards': row['standards'], 
        'grade': row['grade'],
        'learning_targets': row['learning_targets_str'],
        'second_learning_target': row['learning_targets'][1]
    })
    cd_out = cd.invoke({
        'standards': row['standards'], 
        'grade': row['grade'],
        'second_learning_target': row['learning_targets'][1]
    })
    lp_out = lp.invoke({
        'standards': row['standards'], 
        'grade': row['grade'],
        'learning_targets': row['learning_targets_str'],
        'second_learning_target': row['learning_targets'][1], 
    })

    prompts = {
        'run_activity_1': p1_out.text,
        'run_activity_2': p2_out.text,       
        'run_warmup': wu_out.text,
        'run_lesson_synthesis': ls_out.text,
        'run_cool_down': cd_out.text,
        'run_lesson_plan': lp_out.text
    }

    return prompts

def get_last_assistant_message(thread_id, to_skip=0):
    messages = client.beta.threads.messages.list(
                    thread_id=thread_id
                ) 
    assistant_messages = []
    user_count = 0
    l = len(messages.data)
    for msg in reversed(messages.data[:l-to_skip]):
        if msg.role == 'assistant':
            assistant_messages.append(msg.content[0].text.value)
        elif msg.role == 'user':
            user_count += 1
            if user_count == 2:
                break
    assistant_messages = '\n'.join(assistant_messages)
    try:
        cleaned_text = assistant_messages.replace("json\n", "").replace("```", "").replace("\n", "")
        cleaned_text = cleaned_text.replace("\\", "\\\\")
        return json.loads(cleaned_text)
    except:
        return assistant_messages
    

def dict_to_markdown(d:dict, prefix=''):
    md = ''
    for key, value in d.items():
        if key.strip().lower() not in ("lesson_title", "lesson_timeline", "duration", 'questions'):
            # handle key and casing, current key is lowercase and concatenated with underscores, want to replace _ with space and capitalize each word
            # Write the main section title
            md += f"##{prefix} {key.replace('_', ' ').capitalize()}\n"
            if isinstance(value, list):
                # Write sections for list data
                for item in value:
                    md += f"- {item}\n"
            elif isinstance(value, dict):
                print(value.keys())
                # Write sections for dictionary data
                for subkey, subvalue in value.items():
                    if isinstance(subvalue, list):
                        md += f"###{prefix} {subkey}\n"
                        for item in subvalue:
                            md += f"- {item}\n"
                    else:
                        md += f"###{prefix} {subkey}\n"
                        md += f"{subvalue}\n"
            else:
                md += f"{value}\n"
        elif key.strip().lower() == 'lesson_timeline':
            md += f"###{prefix} Lesson Timeline\n"
            for subkey, subvalue in value.items():
                md += f"| {subkey} | {subvalue} |\n"
        elif key.strip().lower() == 'questions':
            md += f"###{prefix} Questions\n"
            for idx, q in enumerate(value):
                md += f"###{prefix} {idx + 1}. {q['question']}\n"
                md += f"question type: {q['question_type']}\n"
                md += f"student response: {q['student_response']}\n"
    return md


def convert_dict_to_lesson(data:dict) -> str:
    # run_lesson_plan
    lesson_plan = data['run_lesson_plan']
    lesson = '# ' + lesson_plan['lesson_title'] + '\n'
    print(lesson)
    try:
        lesson += dict_to_markdown(lesson_plan)
    except Exception as e:
        logging.error(f"Failed to convert lesson plan to markdown. Error: {e}")
        lesson += str(lesson_plan)
    lesson += '---\n'

    # run_warmup
    warmup = data['run_warmup']
    lesson += '## Warm-up\n'   
    try:
        lesson += dict_to_markdown(warmup, "#")
    except Exception as e:
        logging.error(f"Failed to convert warmup to markdown. Error: {e}")
        lesson += str(warmup)
    lesson += '---\n'

    # run_activities
    activity_1 = data['run_activity_1']
    lesson += '## Activity 1 \n'
    try:
        lesson += dict_to_markdown(activity_1, "#")
    except Exception as e:
        logging.error(f"Failed to convert activity 1 to markdown. Error: {e}")
        lesson += str(activity_1)

    activity_2 = data['run_activity_2']
    lesson += '## Activity 2 \n'
    try:
        lesson += dict_to_markdown(activity_2, "#")
    except Exception as e:
        logging.error(f"Failed to convert activity 2 to markdown. Error: {e}")
        lesson += str(activity_2)
    lesson += '---\n'

    # run_lesson_synthesis
    synthesis = data['run_lesson_synthesis']
    lesson += '## Lesson Synthesis\n'
    try:
        lesson += dict_to_markdown(synthesis, "#")
    except Exception as e:
        logging.error(f"Failed to convert lesson synthesis to markdown. Error: {e}")
        lesson += str(synthesis)
    lesson += '---\n'

    # run_cool_down
    cool_down = data['run_cool_down']
    lesson += '## Cool-down\n'
    try:
        lesson += dict_to_markdown(cool_down, "#")
    except Exception as e:
        logging.error(f"Failed to convert cool down to markdown. Error: {e}")
        lesson += str(cool_down)

    return lesson

def add_markdown_format(doc, line):
    if line.startswith("#### "):
        p = doc.add_paragraph(line[5:], style='Heading4')
    elif line.startswith("### "):
        p = doc.add_paragraph(line[4:], style='Heading3')
    elif line.startswith("## "):
        p = doc.add_paragraph(line[3:], style='Heading2')
    elif line.startswith("# "):
        p = doc.add_paragraph(line[2:], style='Heading1')
    elif line.startswith("---"):
        doc.add_page_break()
    else:
        p = doc.add_paragraph(line)
def save_text_to_docx(text:str, row_key:str, prefix='output/'):
    doc = Document()
    # Save the markdown text directly to a docx file
    for line in text.strip().split('\n'):
        add_markdown_format(doc, line)
    docx_file =  row_key + '.docx'
    doc.save(prefix + docx_file)
def extract_learning_targets(thread_id) -> str:
    """
    Extracts and returns content from messages where the role is 'assistant',
    specifically content within ```json ...``` format, as a dictionary.
    assuming one message
    
    Args:
    messages (List): A list of message objects.
    
    Returns:
    dict: The extracted JSON content from the assistant messages.

    """
    messages = client.beta.threads.messages.list(thread_id=thread_id) 
    complete =  ''
    for msg in messages.data:
        role = msg.role
        content = msg.content[0].text.value
        if role == 'assistant':
            complete += content
    return complete 

def clean_json_str(json_str:str)->dict:
    """
    Cleans up the JSON string by removing the ```json ...``` and ```...``` tags.
    
    Args:
    json_str (str): A JSON string.
    
    Returns:
    str: A cleaned JSON string.
    """
    json_str = re.sub(r'```json', '', json_str)
    json_str = re.sub(r'```', '', json_str)
    return json.loads(json_str)

def get_grade_from_id(id:str)->str:
    if id in ('1', '2', '3', '4', '5', '6', '7', '8'):
        grade = f'Grade {id}'
    elif id == 'k':
        grade = 'Kindergarten'
    elif id == 'a1':
        grade = 'Algebra 1'
    elif id == 'a2':
        grade = 'Algebra 2'
    elif id == 'g':
        grade = 'Geometry'
    else:
        print("Invalid grade ID. Please provide a valid grade ID.")
        sys.exit(1)
    return grade

def map_id_to_asst_key(id:str)->str:
    if id == 'k':
        return 'k'
    elif id in ('1', '2'):
        return '1-2'
    elif id in ('3', '4', '5'):
        return '3-5'
    elif id in ('6', '7', '8'):
        return '6-8'
    elif id in ('a1', 'a2', 'g'):
        return 'hs'
    else:
        print("Invalid grade ID. Please provide a valid grade ID.")
        sys.exit(1)

@st.cache_data()
def run_activity_1_function(prompts, asst_lw_id, fu_key='run_activity_1'):
    try: 
        thread_lw = client.beta.threads.create(
            messages = [{'role': 'user','content': prompts[fu_key]}],
        )
        run_lw = client.beta.threads.runs.create(thread_id=thread_lw.id, assistant_id=asst_lw_id)
        run_lw = check_run_status(run_lw, thread_lw.id, fu_key)
        return thread_lw
    except Exception as e:
        st.error(f"Failed to run {fu_key}. Error: {e}")
        st.stop()
@st.cache_data()
def execute_lesson_plan_function(thread_lw_id, fu_key, prompts, asst_lw_id, raw_lesson_plan):
    try:
        to_skip = len_message(thread_lw_id)
        content = prompts[fu_key]
        created_message = client.beta.threads.messages.create(
                            thread_lw_id,
                            role="user",
                            content=content,
                            )
        run = client.beta.threads.runs.create(thread_id=thread_lw.id, assistant_id=asst_lw_id)
        run = check_run_status(run, thread_lw.id, fu_key)
        msg = get_last_assistant_message(thread_lw.id, to_skip=to_skip)

        raw_lesson_plan[fu_key] = msg
        return raw_lesson_plan
    except Exception as e:
        st.error(f"Failed to run {fu_key}. Error: {e}")
        st.stop()
if __name__ == '__main__':
    st.set_page_config(page_title='TX Lesson Generator', page_icon=':books:', layout='wide')
    st.title('2-step TX Lesson Generator')

    st.markdown('''
    This app generates a lesson plan for a given grade and standard in two steps:
    1. Generate learning targets with three inputs: grade level, standard, and standard description.
    2. Generate lesson plan with four inputs: learning targets ( generated from above step ), grade level, standard, and standard description.
    ''')
    
    # start a section 
    st.header('Learning Targets Generator')
    # a few inputs: a dropdown for grade level, standard, and description of the standard
    grade_level = st.selectbox('Select Grade Level', ['k', '1', '2', '3', '4', '5', '6', '7', '8', 'a1', 'a2', 'g'], index=3)
    standard = st.text_input('Enter Breakout Number', placeholder='3.F.viii')
    standard_description = st.text_area('Enter Standard Description (KSS. Expectation. Breakout Description)', height=150,
                                         placeholder='''Number and operations. 
The student applies mathematical process standards to represent and explain fractional units. 
The student is expected to: 
    represent equivalent fractions with denominators of 2, 3, 4, 6, and 8 using a variety of objects and pictorial models, including number lines;
    represent equivalent fractions with denominators of 4 using a variety of pictorial models, including number lines''')

    grade = get_grade_from_id(grade_level)
    asst_key = map_id_to_asst_key(grade_level)
    asst_lg_id = assistant_ids[asst_key]['asst_lg']['id']
    asst_lw_id = assistant_ids[asst_key]['asst_lw']['id']

    # st.write(f'Grade: {grade} Assessment Key: {asst_key}')
    # st.write(f'Standard: {standard}')
    # st.write(f'Standard Description: {standard_description}')
    if not grade_level or not standard or not standard_description:
        st.error("Please ensure all fields (Grade Level, Standard, Standard Description) are filled.")
        st.stop()

    st.write('Generating learning targets...')
    lg_start = time.time()
    thread, run = run_learning_targets(standard, standard_description, grade, asst_lg_id)
    lg_end = time.time()
    st.caption(f'Time taken: {lg_end - lg_start:.2f} seconds')
    # retrieve learning targets
    lts = extract_learning_targets(thread.id)
    lts_ls = clean_json_str(lts)
    # st.write(f'Learning Targets: {lts_ls}, type: {type(lts_ls)}')
    lts_ls_str = lts_ls['lt1'] + '\n' + lts_ls['lt2']
    final_lts = st.text_area('Learning Targets', value=lts_ls_str, height=120)
    st.warning('''Now, review the learning targets and make sure that there are two "I can" statements that are clear and appropriate 
               for the grade level and the standards before proceeding to the next step.''')
    # convert final_lts to a list of two 
    final_lts_ls = final_lts.split('\n')
    if len(final_lts_ls) != 2:
        st.error("Please ensure there are two learning targets; there should be two 'I can' statements separated by a new line.")
        st.stop()
    
    st.header('Lesson Plan Generator')
    if st.button('Generate Lesson Plan'):
        lw_start = time.time()
        raw_lesson_plan = {}
        row = {
            'grade': grade,
            'standards': standard,
            'standard_description': standard_description,
            'learning_targets': final_lts_ls,
            'learning_targets_str': final_lts
        }
        row_key = f'G{grade_level}_Breakout{standard}_New Lesson'
        prompts = create_all_prompts(row)
        
        fu_key = 'run_activity_1'
        thread_lw = run_activity_1_function(prompts, asst_lw_id)
        msg = get_last_assistant_message(thread_lw.id, to_skip=0)
        raw_lesson_plan[fu_key] = msg
        # st.write(f'Activity 1: {msg}')
        
        for fu_key in ('run_activity_2', 'run_warmup', 'run_lesson_synthesis', 'run_cool_down', 'run_lesson_plan'):
            st.spinner(f'Generating {fu_key}...')
            raw_lesson_plan = execute_lesson_plan_function(thread_lw.id, fu_key, prompts, asst_lw_id, raw_lesson_plan)

        # st.write(f'Raw Lesson Plan: {raw_lesson_plan}')
        # st.write(f'{row_key}, {raw_lesson_plan.keys()}')
        # check if output_pieces has the same keys as prompts
        assert set(raw_lesson_plan.keys()) == set(prompts.keys())
        lesson = convert_dict_to_lesson(raw_lesson_plan)
        lw_end = time.time()
        st.caption(f'Time taken: {(lw_end - lw_start)/60:.2f} minutes')
        # save_text_to_docx(lesson, row_key)
        with st.expander('View Lesson Plan'):
            st.write(lesson)
        # create a download button to save the doc into a docx file
        # The code remains unchanged as no specific instructions were provided for rewriting.
        def str_to_doc(text:str):
            doc = Document()
            # Save the markdown text directly to a docx file
            for line in text.strip().split('\n'):
                add_markdown_format(doc, line)
            return doc
    
        doc_to_download = str_to_doc(lesson)
        bio = io.BytesIO()
        doc_to_download.save(bio)
        if doc_to_download:
            st.download_button(
                label="Click here to download",
                data=bio.getvalue(),
                file_name=f"{row_key}.docx",
                mime="docx"
            )


    
